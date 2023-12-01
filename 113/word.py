from docx import Document
from docx.shared import Cm

documento = Document()
documento.add_heading('Titulo do documento',0) #adc titulo

#paragrafo
paragrafo = documento.add_paragraph('Um parágrafo simples ')
paragrafo.add_run('e importante').bold = True
paragrafo.add_run('com palavras em ')
paragrafo.add_run('itálico').italic = True 


#criando outros titulos como se fosse h1,h2,h3
documento.add_heading('Titulo - Nível 1', level = 1)
documento.add_heading('Titulo - Nível 2', level = 2)
documento.add_heading('Titulo - Nível 3', level = 3)

#Estilo de paragrafos
documento.add_paragraph('Formatação "No Spacing "', style='No Spacing')
documento.add_paragraph('Formatação "Heading 1"', style='Heading 1')
documento.add_paragraph('Formatação "Heading 2 "', style='Heading 2')
documento.add_paragraph('Formatação "Heading 3"', style='Heading 3')
documento.add_paragraph('Formatação "Title "', style='Title')
documento.add_paragraph('Formatação "Subtitle "', style='Subtitle')
documento.add_paragraph('Formatação "Quote "', style='Quote')
documento.add_paragraph('Formatação "Intense Quote  "', style='Intense Quote')
documento.add_paragraph('Formatação "List Paragraph "', style='List Paragraph')
documento.add_paragraph('Formatação "List Bullet "', style='List Bullet')
documento.add_paragraph('Formatação "List Number "', style='List Number')

#adicionando fotos
documento.add_picture('teste.png', width=Cm(5.25) )

#criando tabelas
tabela = documento.add_table(rows=3,cols=2)
celula1 = tabela.cell(0,0)
celula1.text = 'Nome'


#quebrar pagina
documento.add_page_break()

#outra maneira de criar tabelas usando tuplas
compras = (
    (3,'101','Uva'),
    (4,'301','Pão'),
    (8,'341','Banana, mamão')
)
tabela_supermercado = documento.add_table(rows=1,cols=3)
cabecalho_table_supermercado = tabela_supermercado.rows[0].cells
cabecalho_table_supermercado[0].text = 'Quantidade'
cabecalho_table_supermercado[1].text = 'Id'
cabecalho_table_supermercado[2].text = 'Descrição'

for quantidade, id,desc in compras:
    linha = tabela_supermercado.add_row().cells
    linha[0].text = str(quantidade)
    linha[1].text = id
    linha[2].text = desc


documento.add_page_break()
documento.add_paragraph("Essa é uma página nova")

documento.save('demo.docx')














